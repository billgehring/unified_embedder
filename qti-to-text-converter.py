import xml.etree.ElementTree as ET
import html
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def find_main_assessment_file(manifest_path):
    tree = ET.parse(manifest_path)
    root = tree.getroot()
    
    ns = {'ns': 'http://www.imsglobal.org/xsd/imsccv1p1/imscp_v1p1'}
    
    resource = root.find(".//ns:resource[@type='imsqti_xmlv1p2']", ns)
    
    if resource is not None:
        file_elem = resource.find("ns:file", ns)
        if file_elem is not None:
            return file_elem.get('href')
    
    return None

def parse_qti(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()

    ns = {'': 'http://www.imsglobal.org/xsd/ims_qtiasiv1p2'}

    questions = []

    for item in root.findall('.//item', ns):
        question = {}
        question['title'] = item.get('title', '')
        
        material = item.find('.//material/mattext', ns)
        if material is not None:
            question['text'] = material.text or ''

        question['choices'] = []
        for response_label in item.findall('.//response_lid/render_choice/response_label', ns):
            choice_text = response_label.find('.//material/mattext', ns)
            if choice_text is not None:
                question['choices'].append({
                    'id': response_label.get('ident'),
                    'text': choice_text.text or ''
                })

        correct_answer = item.find('.//respcondition/conditionvar/varequal', ns)
        if correct_answer is not None:
            question['correct_answer'] = correct_answer.text

        questions.append(question)

    return questions

def clean_html(text):
    # Remove style attributes
    text = re.sub(r' style="[^"]*"', '', text)
    # Convert <p> tags to newlines
    text = re.sub(r'<p.*?>(.*?)</p>', r'\1\n', text, flags=re.DOTALL)
    # Convert <ul> and <ol> to newlines
    text = re.sub(r'<[uo]l.*?>(.*?)</[uo]l>', r'\n\1', text, flags=re.DOTALL)
    # Convert <li> to bullet points
    text = re.sub(r'<li.*?>(.*?)</li>', r'• \1\n', text, flags=re.DOTALL)
    # Remove remaining HTML tags
    text = re.sub(r'<.*?>', '', text)
    # Fix quotation marks
    text = text.replace('"', '"').replace('"', '"')
    # Remove extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    return text

def format_questions_html(questions):
    formatted_html = "<html><body>"
    for i, q in enumerate(questions, 1):
        formatted_html += f"<h3>{i}. {clean_html(q['text'])}</h3>"
        
        formatted_html += "<ol type='a'>"
        for choice in q['choices']:
            correct = " &#10004;" if choice['id'] == q.get('correct_answer') else ""
            formatted_html += f"<li>{clean_html(choice['text'])}{correct}</li>"
        formatted_html += "</ol>"
    
    formatted_html += "</body></html>"
    return formatted_html

def format_questions_markdown(questions):
    formatted_md = ""
    for i, q in enumerate(questions, 1):
        formatted_md += f"## {i}. {clean_html(q['text'])}\n\n"
        
        for j, choice in enumerate(q['choices']):
            letter = chr(97 + j)  # Convert 0, 1, 2, 3 to a, b, c, d
            correct = "*" if choice['id'] == q.get('correct_answer') else ""
            formatted_md += f"{correct}{letter}) {clean_html(choice['text'])}\n"
        
        formatted_md += "\n"
    
    return formatted_md

def format_questions_docx(questions, doc):
    for i, q in enumerate(questions, 1):
        doc.add_heading(f"{i}. {clean_html(q['text'])}", level=2)
        
        for j, choice in enumerate(q['choices']):
            letter = chr(97 + j)  # Convert 0, 1, 2, 3 to a, b, c, d
            p = doc.add_paragraph(f"{letter}) {clean_html(choice['text'])}")
            if choice['id'] == q.get('correct_answer'):
                run = p.add_run(" ✓")
                run.font.name = 'Segoe UI Symbol'
                run.font.size = Pt(12)
        
        doc.add_paragraph()  # Add a blank line between questions

def convert_qti_to_multiple_formats(qti_folder_path, output_folder):
    manifest_path = os.path.join(qti_folder_path, 'imsmanifest.xml')
    
    if not os.path.exists(manifest_path):
        print(f"Error: imsmanifest.xml not found in {qti_folder_path}")
        return
    
    main_assessment_file = find_main_assessment_file(manifest_path)
    
    if main_assessment_file is None:
        print("Error: Unable to find the main assessment file in the manifest.")
        return
    
    full_assessment_path = os.path.join(qti_folder_path, main_assessment_file)
    
    if not os.path.exists(full_assessment_path):
        print(f"Error: Assessment file not found: {full_assessment_path}")
        return
    
    questions = parse_qti(full_assessment_path)
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Generate HTML output
    html_output = format_questions_html(questions)
    with open(os.path.join(output_folder, 'output.html'), 'w', encoding='utf-8') as f:
        f.write(html_output)

    # Generate Markdown output
    md_output = format_questions_markdown(questions)
    with open(os.path.join(output_folder, 'output.md'), 'w', encoding='utf-8') as f:
        f.write(md_output)

    # Generate Word document output
    doc = Document()
    format_questions_docx(questions, doc)
    doc.save(os.path.join(output_folder, 'output.docx'))

    print(f"Conversion complete. Outputs saved in {output_folder}")

# Usage
qti_folder_path = r'/Users/wgehring/University of Michigan Dropbox/William Gehring/Classes/Psych240/2024 SU/Exam 2 Materials/Summer 2024 Week 6 Quiz final version'  # This should be the folder containing imsmanifest.xml
output_folder = r'/Users/wgehring/University of Michigan Dropbox/William Gehring/Classes/Psych240/2024 SU/Exam 2 Materials'
convert_qti_to_multiple_formats(qti_folder_path, output_folder)
